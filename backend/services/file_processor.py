import os
import io
import json
import easyocr
from PIL import Image
from PyPDF2 import PdfReader
from docx import Document
import mailparser as mail_parser
from typing import Dict, Any, Optional
from werkzeug.datastructures import FileStorage

class FileProcessor:
    def __init__(self):
        # Initialize EasyOCR reader (supports multiple languages)
        self.ocr_reader = easyocr.Reader(['en'])  # Initialize with English
    
    def process_file(self, file: FileStorage) -> Optional[Dict[str, Any]]:
        """Process uploaded file and extract data"""
        
        filename = file.filename.lower()
        file_ext = os.path.splitext(filename)[1]
        
        try:
            if file_ext in ['.jpg', '.jpeg', '.png', '.bmp', '.gif']:
                return self._process_image(file)
            elif file_ext == '.pdf':
                return self._process_pdf(file)
            elif file_ext in ['.doc', '.docx']:
                return self._process_word(file)
            elif file_ext in ['.eml', '.msg']:
                return self._process_email(file)
            elif file_ext == '.txt':
                return self._process_text(file)
            elif file_ext == '.json':
                return self._process_json(file)
            else:
                # Try to process as text
                return self._process_text(file)
        except Exception as e:
            print(f"Error processing file: {e}")
            return None
    
    def _process_image(self, file: FileStorage) -> Dict[str, Any]:
        """Extract text from image using OCR"""
        
        image = Image.open(file.stream)
        
        # Perform OCR using EasyOCR
        import numpy as np
        image_array = np.array(image)
        ocr_results = self.ocr_reader.readtext(image_array)
        
        # Extract text from OCR results
        text = ' '.join([result[1] for result in ocr_results])
        
        # Also save image data for vision API
        file.stream.seek(0)
        image_data = file.stream.read()
        
        return {
            'type': 'image',
            'text': text,
            'image_data': image_data,
            'filename': file.filename
        }
    
    def _process_pdf(self, file: FileStorage) -> Dict[str, Any]:
        """Extract text from PDF"""
        
        text_content = []
        pdf_reader = PdfReader(file.stream)
        
        for page in pdf_reader.pages:
            text_content.append(page.extract_text())
        
        return {
            'type': 'pdf',
            'text': '\n'.join(text_content),
            'num_pages': len(pdf_reader.pages),
            'filename': file.filename
        }
    
    def _process_word(self, file: FileStorage) -> Dict[str, Any]:
        """Extract text from Word document"""
        
        doc = Document(file.stream)
        text_content = []
        
        for paragraph in doc.paragraphs:
            text_content.append(paragraph.text)
        
        # Extract tables
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        
        return {
            'type': 'word',
            'text': '\n'.join(text_content),
            'tables': tables,
            'filename': file.filename
        }
    
    def _process_email(self, file: FileStorage) -> Dict[str, Any]:
        """Extract content from email files"""
        
        mail = mail_parser.parse_from_bytes(file.stream.read())
        
        # Extract attachments info
        attachments = []
        for attachment in mail.attachments:
            attachments.append({
                'filename': attachment['filename'],
                'type': attachment.get('content-type', 'unknown')
            })
        
        return {
            'type': 'email',
            'subject': mail.subject,
            'from': mail.from_,
            'to': mail.to,
            'date': str(mail.date),
            'text': mail.text_plain[0] if mail.text_plain else mail.text_html[0] if mail.text_html else '',
            'attachments': attachments,
            'filename': file.filename
        }
    
    def _process_text(self, file: FileStorage) -> Dict[str, Any]:
        """Process plain text file"""
        
        content = file.stream.read().decode('utf-8', errors='ignore')
        
        return {
            'type': 'text',
            'text': content,
            'filename': file.filename
        }
    
    def _process_json(self, file: FileStorage) -> Dict[str, Any]:
        """Process JSON file"""
        
        content = json.load(file.stream)
        
        return {
            'type': 'json',
            'data': content,
            'text': json.dumps(content, indent=2),
            'filename': file.filename
        }
    
    def extract_product_info(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Extract structured product information from processed data"""
        
        text = data.get('text', '')
        
        # Basic extraction patterns (can be enhanced with NLP)
        info = {
            'specifications': {},
            'features': [],
            'model_numbers': []
        }
        
        # Extract common HVAC specifications
        import re
        
        # Tonnage
        tonnage_match = re.search(r'(\d+(?:\.\d+)?)\s*ton', text, re.IGNORECASE)
        if tonnage_match:
            info['specifications']['tonnage'] = float(tonnage_match.group(1))
        
        # SEER rating
        seer_match = re.search(r'(\d+(?:\.\d+)?)\s*SEER', text, re.IGNORECASE)
        if seer_match:
            info['specifications']['seer'] = float(seer_match.group(1))
        
        # Model numbers (alphanumeric patterns)
        model_patterns = re.findall(r'\b[A-Z]{2,4}[-\s]?\d{2,6}[A-Z]?\b', text)
        info['model_numbers'] = model_patterns
        
        return info